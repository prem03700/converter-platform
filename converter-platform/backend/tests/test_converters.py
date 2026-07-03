"""
Converter tests. These exercise the REAL engines (Pillow, PyMuPDF,
LibreOffice via subprocess, FFmpeg via subprocess, Tesseract) — nothing
here is mocked. Run with: pytest tests/test_converters.py -v
"""
import io
import json
import math
import struct
import tarfile
import wave
import zipfile

import pytest
from docx import Document

from app.converters.registry import get_converter


def test_text_family_txt_to_html():
    conv = get_converter("txt", "html")
    res = conv.convert(b"Hello world\nSecond line", "txt", "html")
    assert res.output_mime_type == "text/html"
    assert b"Hello world" in res.data


def test_json_to_yaml_to_xml_round_trip():
    payload = json.dumps({"name": "test", "values": [1, 2, 3]}).encode()
    yaml_conv = get_converter("json", "yaml")
    yaml_result = yaml_conv.convert(payload, "json", "yaml")
    assert b"name: test" in yaml_result.data

    xml_conv = get_converter("yaml", "xml")
    xml_result = xml_conv.convert(yaml_result.data, "yaml", "xml")
    assert b"<name>test</name>" in xml_result.data


def test_image_png_to_webp():
    from PIL import Image

    img = Image.new("RGB", (50, 50), color=(10, 20, 30))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    conv = get_converter("png", "webp")
    res = conv.convert(buf.getvalue(), "png", "webp")
    assert res.output_mime_type == "image/webp"
    assert len(res.data) > 0
    Image.open(io.BytesIO(res.data)).verify()


def test_svg_to_png_rasterization():
    svg = (
        b'<svg xmlns="http://www.w3.org/2000/svg" width="100" height="100">'
        b'<circle cx="50" cy="50" r="40" fill="red"/></svg>'
    )
    conv = get_converter("svg", "png")
    res = conv.convert(svg, "svg", "png")
    assert res.output_mime_type == "image/png"
    assert len(res.data) > 100


def test_docx_to_pdf_via_libreoffice():
    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is a paragraph created for conversion testing.")
    buf = io.BytesIO()
    doc.save(buf)

    conv = get_converter("docx", "pdf")
    res = conv.convert(buf.getvalue(), "docx", "pdf")
    assert res.output_mime_type == "application/pdf"
    assert res.data[:4] == b"%PDF"


def test_pdf_to_txt_extraction():
    doc = Document()
    doc.add_paragraph("Round trip extraction test.")
    docx_buf = io.BytesIO()
    doc.save(docx_buf)

    pdf_bytes = get_converter("docx", "pdf").convert(docx_buf.getvalue(), "docx", "pdf").data
    txt_result = get_converter("pdf", "txt").convert(pdf_bytes, "pdf", "txt")
    assert "Round trip extraction test." in txt_result.data.decode()


def test_audio_wav_to_mp3_via_ffmpeg():
    buf = io.BytesIO()
    with wave.open(buf, "wb") as w:
        w.setnchannels(1)
        w.setsampwidth(2)
        w.setframerate(8000)
        frames = b"".join(
            struct.pack("<h", int(3000 * math.sin(2 * math.pi * 440 * t / 8000)))
            for t in range(8000)
        )
        w.writeframes(frames)

    conv = get_converter("wav", "mp3")
    res = conv.convert(buf.getvalue(), "wav", "mp3")
    assert res.output_mime_type == "audio/mpeg"
    assert len(res.data) > 0


def test_zip_to_tar_archive_conversion():
    zbuf = io.BytesIO()
    with zipfile.ZipFile(zbuf, "w") as zf:
        zf.writestr("hello.txt", "hello from inside the zip")

    conv = get_converter("zip", "tar")
    res = conv.convert(zbuf.getvalue(), "zip", "tar")
    with tarfile.open(fileobj=io.BytesIO(res.data)) as tf:
        assert "./hello.txt" in tf.getnames()


def test_python_code_to_highlighted_html():
    conv = get_converter("py", "html")
    res = conv.convert(b"def hello():\n    print('hi')\n", "py", "html")
    assert res.output_mime_type == "text/html"
    assert b"hello" in res.data


def test_unsupported_pair_raises_conversion_error():
    from app.converters.base import ConversionError

    with pytest.raises(ConversionError):
        get_converter("mp3", "docx")


def test_ocr_image_to_text_runs_locally():
    from PIL import Image, ImageDraw

    from app.services import ai_service

    img = Image.new("RGB", (300, 80), color="white")
    d = ImageDraw.Draw(img)
    d.text((10, 30), "Hello OCR Test", fill="black")
    buf = io.BytesIO()
    img.save(buf, format="PNG")

    text = ai_service.image_to_text(buf.getvalue())
    assert "ocr" in text.lower()
